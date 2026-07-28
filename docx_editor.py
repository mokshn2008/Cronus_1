"""
Lightweight DOCX editing for Cronus.

Scope, stated honestly: this handles paragraph-level TEXT edits (fixing
wording, grammar, rewriting sections) while preserving each paragraph's
basic style (based on its first run). It does NOT handle complex
formatting fidelity, tables, images, or tracked changes -- that's a much
larger undertaking (see /mnt/skills/public/docx/SKILL.md's XML-level
approach if you ever need that level of fidelity).
"""

import json
from pathlib import Path
from typing import List, Optional
from docx import Document


def extract_paragraphs(docx_path: str) -> List[str]:
    doc = Document(docx_path)
    return [p.text for p in doc.paragraphs]


def apply_paragraph_edits(original_path: str, edited_paragraphs: List[str], output_path: str) -> str:
    """
    Writes edited_paragraphs back into a copy of the original document,
    preserving each paragraph's basic formatting (based on its first run).
    edited_paragraphs must be the same length as the original paragraph list --
    if a paragraph is unchanged, pass its original text back unchanged.
    """
    doc = Document(original_path)
    original_paragraphs = doc.paragraphs

    if len(edited_paragraphs) != len(original_paragraphs):
        raise ValueError(
            f"Edited paragraph count ({len(edited_paragraphs)}) doesn't match "
            f"original ({len(original_paragraphs)}) -- refusing to write, "
            f"since this would misalign edits with the wrong paragraphs."
        )

    for para, new_text in zip(original_paragraphs, edited_paragraphs):
        if para.text == new_text:
            continue  # leave unchanged paragraphs completely untouched

        if para.runs:
            # Keep the first run's formatting, put all new text there,
            # clear any additional runs so we don't duplicate text.
            para.runs[0].text = new_text
            for run in para.runs[1:]:
                run.text = ""
        else:
            para.add_run(new_text)

    doc.save(output_path)
    return output_path


def build_edit_prompt(paragraphs: list[str], instruction: str) -> str:
    numbered = "\n".join(f"[{i}] {p}" for i, p in enumerate(paragraphs))
    return f"""Here is a document broken into numbered paragraphs:

{numbered}

Instruction: {instruction}

Return ONLY a JSON array of strings, one per paragraph, in the SAME ORDER,
same length as the input ({len(paragraphs)} items). For any paragraph you
are not changing, return its original text exactly unchanged. No other
text, no markdown formatting, just the raw JSON array."""


def parse_edit_response(raw_response: str, expected_count: int) -> Optional[List[str]]:
    """Returns None if parsing fails or count doesn't match -- caller should
    fall back to returning the original document unedited rather than guess."""
    cleaned = raw_response.strip()
    if cleaned.startswith("```"):
        cleaned = cleaned.split("```")[1]
        if cleaned.startswith("json"):
            cleaned = cleaned[4:]
    try:
        parsed = json.loads(cleaned)
        if isinstance(parsed, list) and len(parsed) == expected_count:
            return parsed
    except json.JSONDecodeError:
        pass
    return None
